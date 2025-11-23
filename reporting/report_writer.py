from docx import Document
import pandas as pd

def df_to_word_table(doc: Document, df: pd.DataFrame, title: str = None):
    """
    Insert a Pandas DataFrame as a table into a Word document.
    Handles MultiIndex columns (e.g., from describe()).
    
    Parameters:
        doc (Document): Word document object
        df (pd.DataFrame): DataFrame to insert
        title (str): Optional title for the table
    """
    if title:
        doc.add_heading(title, level=3)
    
    if df.empty:
        doc.add_paragraph("No data available.")
        return

    # Flatten MultiIndex columns if present
    if isinstance(df.columns, pd.MultiIndex):
        df.columns = [' | '.join([str(i) for i in col]).strip() for col in df.columns]

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'LightList'  # Optional style

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    # Add rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)


def generate_report(summary_df: pd.DataFrame, output_path: str, report_title: str = "Automated Data Summary Report"):
    """
    Generate a Word report from a summary DataFrame.
    
    Parameters:
        summary_df (pd.DataFrame): Summary DataFrame (from get_user_summary)
        output_path (str): Path to save the Word report
        report_title (str): Title for the report
    """
    doc = Document()

    # Introduction
    doc.add_heading(report_title, level=1)
    doc.add_paragraph("This report provides an automated summary of the uploaded dataset.")

    # Background
    doc.add_heading("Background", level=2)
    doc.add_paragraph(
        "The purpose of this exercise is to generate analytics reports "
        "from data sources using user-selected grouping and summary options."
    )

    # Scope
    doc.add_heading("Scope", level=2)
    doc.add_paragraph(
        "This report summarizes data based on user-selected columns and "
        "summary type. Only the specified columns and groupings are included."
    )

    # Findings
    doc.add_heading("Findings", level=2)
    doc.add_paragraph(f"Total Rows in summary: {len(summary_df)}")
    doc.add_paragraph(f"Total Columns in summary: {len(summary_df.columns)}")

    # Insert summary table
    df_to_word_table(doc, summary_df, title="User-Selected Summary")

    # Conclusion
    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph("The dataset has been successfully summarized and processed.")

    # Save document
    doc.save(output_path)
    print(f"Report generated: {output_path}")
